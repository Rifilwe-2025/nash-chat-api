"""Word documents, via ``python-docx`` (spec §5.2.3).

Headings are preserved as Markdown ``#`` levels rather than flattened into the body. That matters
downstream twice: Tier 1 injects the whole document, and a model reading a wall of undifferentiated
paragraphs loses the structure that says which policy a sentence belongs to; Tier 2 (Phase 6)
selects sections, and a heading is what names one.

Tables are emitted as pipe-separated rows — the same reasoning as the CSV extractor, but without
inventing sentences from columns whose meaning is not known here.
"""

from __future__ import annotations

import io

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.modules.knowledge_base.internal.extractors.base import (
    ExtractedContent,
    ExtractionError,
    ExtractionResult,
)

MAX_HEADING_LEVEL = 6


def _heading_level(paragraph: Paragraph) -> int | None:
    """``Heading 2`` → 2. Returns ``None`` for body text."""
    style = (paragraph.style.name or "") if paragraph.style is not None else ""
    if not style.startswith("Heading"):
        return None
    _, _, tail = style.partition(" ")
    if not tail.strip().isdigit():
        return 1  # "Title" and unnumbered heading styles read as the top level
    return min(int(tail.strip()), MAX_HEADING_LEVEL)


def _render_table(table: Table) -> str:
    rows = [
        " | ".join(cell.text.strip() for cell in row.cells)
        for row in table.rows
        if any(cell.text.strip() for cell in row.cells)
    ]
    return "\n".join(rows)


class DocxExtractor:
    def __init__(self, max_heading_level: int = MAX_HEADING_LEVEL) -> None:
        self._max_heading_level = max_heading_level

    async def extract(self, content: ExtractedContent) -> ExtractionResult:
        try:
            document = docx.Document(io.BytesIO(content.data))
        except Exception as exc:  # python-docx raises a package-specific error for bad zips
            raise ExtractionError("The file could not be read as a Word document.") from exc

        blocks: list[str] = []
        headings = 0

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            level = _heading_level(paragraph)
            if level is None:
                blocks.append(text)
            else:
                headings += 1
                blocks.append(f"{'#' * min(level, self._max_heading_level)} {text}")

        for table in document.tables:
            rendered = _render_table(table)
            if rendered:
                blocks.append(rendered)

        if not blocks:
            raise ExtractionError("The document contains no readable text.")

        text = "\n\n".join(blocks)
        return ExtractionResult(
            text=text,
            metadata={
                "format": "docx",
                "headings": headings,
                "paragraphs": len(blocks) - headings,
                "tables": len(document.tables),
                "characters": len(text),
            },
        )
